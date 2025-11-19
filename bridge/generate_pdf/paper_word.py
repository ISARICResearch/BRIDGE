import io
import re
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

MAX_CHOICES_SHOWN = 10


def pick_col(df, names):
    for n in names:
        if n in df.columns:
            return n
    return None

def split_choices(raw):
    if pd.isna(raw) or not str(raw).strip():
        return []
    parts = re.split(r'\s*\|\s*|\r?\n|;\s*', str(raw).strip())
    return [p.strip() for p in parts if p.strip()]

def clean_choice_label(part):
    p = str(part).strip()
    for sep in [',', ':', '=']:
        if sep in p:
            return p.split(sep, 1)[1].strip()
    m = re.split(r'^\s*\d+[\)\.:-]\s*', p)
    return m[-1].strip() if m else p

def is_date(field_type, validation):
    ft = str(field_type or "").lower()
    val = str(validation or "").lower()
    return (
        "date" in ft or "datetime" in ft or
        "date_dmy" in val or "date_mdy" in val or "date_ymd" in val or "date" in val
    )

RE_88 = re.compile(
    r"(?:=\s*['\"]?\s*88\s*['\"]?)"
    r"|(?:\(\s*88\s*\)\s*=\s*['\"]?\s*1\s*['\"]?)"
    r"|(?:\b88\s*=\s*['\"]?\s*1\s*['\"]?)",
    re.I,
)

def branching_other(txt):
    return bool(RE_88.search(str(txt or "")))

def format_answer(ftype, choices_raw, is_date_field):
    ft = str(ftype or "").lower().strip()
    if is_date_field:
        return ["[DD-MM-YYY]"]
    choices = [clean_choice_label(x) for x in split_choices(choices_raw)]
    if ft in ("radio", "dropdown", "select"):
        if choices:
            shown = choices[:MAX_CHOICES_SHOWN]
            return [" ".join(f"○ {c}" for c in shown) + (" ...↓" if len(choices) > MAX_CHOICES_SHOWN else "")]
        return ["○ ________"]
    if ft in ("checkbox", "check box", "checks"):
        if choices:
            shown = choices[:MAX_CHOICES_SHOWN]
            return [" ".join(f"☐ {c}" for c in shown) + (" ...↓" if len(choices) > MAX_CHOICES_SHOWN else "")]
        return ["☐ ________"]
    if ft in ("text", "notes", "textarea"):
        return ["__________"]
    if ft in ("calc", "calculation"):
        return ["Calculated value (read-only)"]
    if choices:
        shown = choices[:MAX_CHOICES_SHOWN]
        return [" ".join(f"- {c}" for c in shown) + (" ...↓" if len(choices) > MAX_CHOICES_SHOWN else "")]
    return ["__________"]


def df_to_word(df: pd.DataFrame) -> bytes:
    """
    Genera un .docx tipo 'paper-like' a partir de un DataFrame REDCap-style
    y devuelve los bytes del documento.
    """
    
    df = df.copy()
    df = df.fillna("")

    c_form = pick_col(df, ["Form Name", "form_name", "Form"])
    c_section = pick_col(df, ["Section Header", "section_header", "Section"])
    c_label = pick_col(df, ["Field Label", "field_label"]) or df.columns[0]
    c_var = pick_col(df, ["Variable / Field Name", "variable / field name", "varname"]) or df.columns[0]
    c_type = pick_col(df, ["Field Type", "field_type"])
    c_choices = pick_col(df, ["Choices, Calculations, OR Slider Labels", "Choices"])
    c_val = pick_col(df, ["Text Validation Type OR Show Slider Number",
                          "Text Validation Type", "text_validation_type"])
    c_branch = pick_col(df, ["Branching Logic (Show field only if...)",
                             "branching_logic", "Logic"])

    # skip descriptive
    if c_type in df.columns:
        df = df[df[c_type].str.lower() != "descriptive"]
        df = df[~df[c_label].str.contains('>')]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    current_form = None
    current_section = None
    table = None

    for _, row in df.iterrows():
        form = str(row.get(c_form, "") or "").strip()

        raw_section = str(row.get(c_section, "") or "").strip()
        if raw_section:
            section = raw_section
        else:
            section = current_section or ""

        if form != current_form:
            current_form = form
            current_section = None
            if form:
                doc.add_paragraph(form).style = doc.styles["Heading 1"]

        if section != current_section:
            current_section = section
            if section:
                doc.add_paragraph(section).style = doc.styles["Heading 2"]

            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Question"
            hdr[1].text = "Answer"
            hdr[2].text = "Notes & justifications"

        question = str(row.get(c_label, "") or row.get(c_var, ""))

        if branching_other(row.get(c_branch, "")) and "(If other)" not in question:
            question += " (If other)"

        ftype = row.get(c_type, "")
        val = row.get(c_val, "")
        isdate = is_date(ftype, val)
        choices = row.get(c_choices, "")

        tr = table.add_row()
        q_cell, a_cell, c_cell = tr.cells
        q_par = q_cell.add_paragraph()
        runq = q_par.add_run(question)
        runq.font.bold = True
        runq.font.size = Pt(9)

        for line in format_answer(ftype, choices, isdate):
            a_cell.add_paragraph(line)

  
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
